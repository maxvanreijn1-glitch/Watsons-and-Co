"""
Generate Watsons & Co legal compliance Word documents.
Run: python3 generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

COMPANY   = "Watsons & Co"
ADDRESS   = "14 Grey Street, Newcastle upon Tyne, NE1 6AE"
PHONE     = "+44 (0) 191 946 0123"
EMAIL     = "hello@watsonsandco.com"
WEBSITE   = "www.watsonsandco.com"
REG_NO    = "09876543"           # England & Wales company number (placeholder)
ICO_REG   = "ZA123456"          # ICO registration number (placeholder)
TPO_NO    = "O-123456"          # The Property Ombudsman number (placeholder)
HMRC_MLR  = "XMWAT00000123456"  # HMRC MLR registration (placeholder)
CMP_PROV  = "Propertymark"      # Client Money Protection provider
YEAR      = "2026"
DATE      = "15 May 2026"

BLACK  = RGBColor(0x0a, 0x0a, 0x0a)
MUTED  = RGBColor(0x88, 0x84, 0x80)
ACCENT = RGBColor(0x1a, 0x3a, 0x2a)


# ── Helpers ──────────────────────────────────────────────────────────────────

def new_doc(title):
    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)

    # Default Normal style
    style = doc.styles['Normal']
    font  = style.font
    font.name  = 'Calibri'
    font.size  = Pt(11)
    font.color.rgb = BLACK

    # Cover header
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COMPANY.upper())
    run.font.name  = 'Garamond'
    run.font.size  = Pt(22)
    run.font.color.rgb = BLACK
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Estate Agents  ·  Newcastle upon Tyne")
    sr.font.size  = Pt(10)
    sr.font.color.rgb = MUTED

    doc.add_paragraph()
    hr = doc.add_paragraph()
    set_para_border(hr)

    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.font.name  = 'Garamond'
    tr.font.size  = Pt(18)
    tr.font.color.rgb = BLACK

    doc.add_paragraph()
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"Last updated: {DATE}")
    dr.font.size  = Pt(9)
    dr.font.color.rgb = MUTED

    doc.add_page_break()
    return doc


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Garamond'
    run.font.size  = Pt(15)
    run.font.bold  = True
    run.font.color.rgb = BLACK
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = 'Calibri'
    run.font.size  = Pt(11.5)
    run.font.bold  = True
    run.font.color.rgb = BLACK
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(3)
    return p


def info_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    lb = p.add_run(f"{label}:  ")
    lb.font.bold = True
    lb.font.size = Pt(11)
    vr = p.add_run(value)
    vr.font.size = Pt(11)
    vr.font.color.rgb = BLACK
    return p


def footer_note(doc):
    doc.add_paragraph()
    hr = doc.add_paragraph()
    set_para_border(hr)
    doc.add_paragraph()
    fn = doc.add_paragraph(
        f"{COMPANY} · {ADDRESS} · {PHONE} · {EMAIL}\n"
        f"Registered in England & Wales · Company No. {REG_NO} · "
        f"Authorised and regulated by The Property Ombudsman (Ref: {TPO_NO})"
    )
    for run in fn.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_para_border(para):
    """Add a thin bottom border to a paragraph."""
    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C4C0B8')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 1. Privacy Policy ─────────────────────────────────────────────────────────

def make_privacy_policy():
    doc = new_doc("Privacy Policy")

    h1(doc, "1.  Introduction")
    body(doc,
        f"{COMPANY} ('we', 'us', 'our') is committed to protecting and respecting your privacy. "
        f"This Privacy Policy explains how we collect, use, disclose, and safeguard your personal "
        f"information when you visit {WEBSITE}, use our services, or engage with us in connection "
        f"with the purchase, sale, or letting of residential property."
    )
    body(doc,
        f"We are registered with the Information Commissioner's Office (ICO) under registration "
        f"number {ICO_REG}. Our registered office is {ADDRESS}."
    )

    h1(doc, "2.  Data Controller")
    info_row(doc, "Company name", COMPANY)
    info_row(doc, "Registered address", ADDRESS)
    info_row(doc, "Company number", REG_NO)
    info_row(doc, "Email", EMAIL)
    info_row(doc, "Telephone", PHONE)
    info_row(doc, "ICO registration", ICO_REG)

    h1(doc, "3.  Personal Data We Collect")
    h2(doc, "3.1  Information you provide to us")
    for item in [
        "Full name, address, and contact details (telephone and email)",
        "Financial information, including proof of funds and mortgage-in-principle documentation",
        "Identity and address verification documents (passport, driving licence, utility bills) for Anti-Money Laundering compliance",
        "Information about your property requirements and preferences",
        "Details of properties you own or are seeking to purchase or let",
        "Correspondence and communications with our team",
    ]:
        bullet(doc, item)

    h2(doc, "3.2  Information we collect automatically")
    for item in [
        "IP address and browser type when you visit our website",
        "Pages visited, time spent, and referring URLs",
        "Cookie data (see our Cookie Policy for full details)",
    ]:
        bullet(doc, item)

    h1(doc, "4.  How We Use Your Personal Data")
    body(doc, "We process your personal data for the following purposes and on the following legal bases under UK GDPR:")

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, t in enumerate(["Purpose", "Legal Basis", "Details"]):
        r = hdr[i].paragraphs[0].add_run(t)
        r.font.bold = True; r.font.size = Pt(10)

    rows_data = [
        ("Providing estate agency services", "Contract", "To carry out the purchase, sale, or letting of property on your behalf."),
        ("Client onboarding and AML checks", "Legal obligation", "We are required by the Money Laundering Regulations 2017 to verify your identity."),
        ("Marketing communications", "Consent / Legitimate interest", "To send you information about properties and services where you have consented or we have a legitimate interest."),
        ("Website operation", "Legitimate interest", "To ensure our website functions correctly and to improve user experience."),
        ("Legal compliance", "Legal obligation", "To comply with applicable laws and regulatory requirements."),
        ("Dispute resolution", "Legitimate interest", "To investigate and resolve complaints or disputes."),
    ]
    for rd in rows_data:
        row = table.add_row().cells
        for i, v in enumerate(rd):
            p = row[i].paragraphs[0]
            p.add_run(v).font.size = Pt(10)

    doc.add_paragraph()

    h1(doc, "5.  Data Sharing")
    body(doc, "We may share your personal data with the following categories of third parties:")
    for item in [
        "Solicitors and conveyancers acting in the transaction",
        "Mortgage brokers and financial advisers with your consent",
        "Surveyors and valuers",
        "HMRC in fulfilment of our Anti-Money Laundering obligations",
        "The Property Ombudsman in connection with any formal complaint",
        "IT service providers and website hosting companies under appropriate data processing agreements",
        "Law enforcement or regulatory bodies where required by law",
    ]:
        bullet(doc, item)
    body(doc, "We do not sell your personal data to third parties.")

    h1(doc, "6.  Data Retention")
    body(doc,
        "We retain personal data only for as long as necessary for the purposes set out in this "
        "policy, and in accordance with our legal obligations. Our standard retention periods are:"
    )
    for item in [
        "Client files and transaction records: 7 years from completion of the transaction",
        "Anti-Money Laundering records: 5 years from the end of the business relationship",
        "Marketing preferences: until you withdraw consent",
        "Website analytics data: 26 months",
        "Correspondence: 3 years from the date of last communication",
    ]:
        bullet(doc, item)

    h1(doc, "7.  Your Rights")
    body(doc, "Under UK GDPR, you have the following rights in relation to your personal data:")
    rights = [
        ("Right of access", "You may request a copy of the personal data we hold about you."),
        ("Right to rectification", "You may ask us to correct inaccurate or incomplete data."),
        ("Right to erasure", "You may ask us to delete your data where there is no compelling reason for its continued processing."),
        ("Right to restrict processing", "You may ask us to restrict how we use your data in certain circumstances."),
        ("Right to data portability", "You may ask us to provide your data in a machine-readable format."),
        ("Right to object", "You may object to processing based on legitimate interests or for direct marketing purposes."),
        ("Right to withdraw consent", "Where processing is based on consent, you may withdraw it at any time."),
    ]
    for title, desc in rights:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{title}: ").font.bold = True
        p.add_run(desc)

    h1(doc, "8.  International Transfers")
    body(doc,
        "We do not routinely transfer personal data outside the UK. Where any transfer is necessary, "
        "we will ensure appropriate safeguards are in place in accordance with UK GDPR."
    )

    h1(doc, "9.  Cookies")
    body(doc,
        "Our website uses cookies. Please refer to our separate Cookie Policy available at "
        f"{WEBSITE}/cookies for full details."
    )

    h1(doc, "10.  How to Exercise Your Rights")
    body(doc,
        f"To exercise any of your rights, or if you have any questions about this Privacy Policy, "
        f"please contact our Data Protection Lead:"
    )
    info_row(doc, "Email", EMAIL)
    info_row(doc, "Post", f"Data Protection Lead, {COMPANY}, {ADDRESS}")
    body(doc,
        "We will respond to all legitimate requests within one calendar month. If your request is "
        "particularly complex, we may extend this period by a further two months and will notify you accordingly."
    )

    h1(doc, "11.  Right to Complain")
    body(doc,
        "If you are unhappy with how we have handled your personal data, you have the right to "
        "lodge a complaint with the Information Commissioner's Office (ICO):"
    )
    info_row(doc, "Website", "www.ico.org.uk")
    info_row(doc, "Helpline", "0303 123 1113")
    info_row(doc, "Post", "Information Commissioner's Office, Wycliffe House, Water Lane, Wilmslow, SK9 5AF")

    h1(doc, "12.  Changes to This Policy")
    body(doc,
        "We may update this Privacy Policy from time to time. The current version will always be "
        f"published on our website at {WEBSITE}. This policy was last updated on {DATE}."
    )

    footer_note(doc)
    doc.save("Privacy_Policy.docx")
    print("✓  Privacy_Policy.docx")


# ── 2. Terms & Conditions ─────────────────────────────────────────────────────

def make_terms():
    doc = new_doc("Terms & Conditions")

    h1(doc, "1.  Introduction")
    body(doc,
        f"These Terms and Conditions govern your use of the {COMPANY} website at {WEBSITE} "
        f"('the Website') and the services provided by {COMPANY}, a company registered in England "
        f"and Wales (Company No. {REG_NO}) whose registered office is at {ADDRESS} ('we', 'us', 'our')."
    )
    body(doc,
        "By accessing our Website or using our services, you agree to be bound by these Terms. "
        "Please read them carefully. If you do not agree, you must not use our Website or services."
    )

    h1(doc, "2.  Our Services")
    body(doc,
        f"{COMPANY} is a residential estate agency operating in Newcastle upon Tyne and Northumberland. "
        f"We provide residential sales, lettings, property management, valuation, and investment advisory services. "
        f"We are members of The Property Ombudsman (Ref: {TPO_NO})."
    )

    h1(doc, "3.  Website Use")
    h2(doc, "3.1  Permitted use")
    body(doc,
        "You may use the Website for lawful purposes only. You must not use it in any way that "
        "breaches applicable law, infringes the rights of others, or transmits unsolicited commercial communications."
    )
    h2(doc, "3.2  Accuracy of information")
    body(doc,
        "While we endeavour to keep the information on the Website accurate and up to date, property "
        "details, prices, and availability are subject to change without notice. Property descriptions "
        "are provided in good faith and do not constitute a representation or warranty."
    )
    h2(doc, "3.3  Intellectual property")
    body(doc,
        f"All content on the Website — including text, photography, graphics, and design — is owned by "
        f"or licensed to {COMPANY} and is protected by copyright and other intellectual property laws. "
        f"You may not reproduce, distribute, or create derivative works without our prior written consent."
    )

    h1(doc, "4.  Estate Agency Services — Terms of Engagement")
    h2(doc, "4.1  Instruction to sell")
    body(doc,
        "By instructing us to market your property for sale, you authorise us to act as your agent "
        "and to market the property on your behalf. Our fee, terms of sole agency or joint agency, "
        "and notice period will be set out in our Letter of Engagement. Fees become payable upon "
        "exchange of contracts (or as otherwise agreed in writing)."
    )
    h2(doc, "4.2  Instruction to let")
    body(doc,
        "By instructing us to find a tenant for your property, you authorise us to act as your lettings "
        "agent. Our fees, services included, and terms will be set out in our Lettings Agreement. "
        "Please refer to our Fees Schedule for a full breakdown of charges."
    )
    h2(doc, "4.3  Consumer Protection from Unfair Trading Regulations 2008")
    body(doc,
        "We comply with the Consumer Protection from Unfair Trading Regulations 2008. All material "
        "information about a property will be disclosed accurately and in a timely manner."
    )
    h2(doc, "4.4  Estate Agents Act 1979")
    body(doc,
        "We operate in accordance with the Estate Agents Act 1979 and keep full records of all "
        "instructions, offers, and transactions. We will disclose any personal interest in a property."
    )

    h1(doc, "5.  Anti-Money Laundering")
    body(doc,
        f"{COMPANY} is registered with HMRC for Anti-Money Laundering supervision (Ref: {HMRC_MLR}). "
        f"We are required by law to verify the identity of all vendors, purchasers, landlords, and tenants "
        f"before proceeding with a transaction. We may request identification documents and, in some cases, "
        f"proof of funds. We may be required to report suspicious activity to the National Crime Agency "
        f"without informing the subject of any such report."
    )

    h1(doc, "6.  Client Money")
    body(doc,
        f"Where we hold client money in connection with lettings and management services, we are members "
        f"of the {CMP_PROV} Client Money Protection scheme. Our client money is held in a separate, "
        f"designated client account and is not mixed with our own funds."
    )

    h1(doc, "7.  Limitation of Liability")
    body(doc,
        "To the fullest extent permitted by law, we exclude all liability for loss or damage — whether "
        "direct, indirect, or consequential — arising from your use of the Website or our services, "
        "except where such liability cannot be excluded by law. Nothing in these Terms excludes or limits "
        "our liability for death or personal injury caused by negligence, fraud, or fraudulent misrepresentation."
    )

    h1(doc, "8.  Third-Party Links")
    body(doc,
        "Our Website may contain links to third-party websites. We have no control over the content of "
        "those sites and accept no responsibility for them or for any loss or damage arising from your use of them."
    )

    h1(doc, "9.  Complaints")
    body(doc,
        "If you are unhappy with any aspect of our service, please refer to our Complaints Procedure "
        f"available at {WEBSITE}/complaints or contact us at {EMAIL}."
    )

    h1(doc, "10.  Governing Law")
    body(doc,
        "These Terms are governed by and construed in accordance with the law of England and Wales. "
        "Any disputes arising under these Terms shall be subject to the exclusive jurisdiction of "
        "the courts of England and Wales."
    )

    h1(doc, "11.  Changes to These Terms")
    body(doc,
        f"We may amend these Terms at any time by posting updated Terms on the Website. "
        f"Continued use of the Website following any changes constitutes acceptance of the new Terms. "
        f"These Terms were last updated on {DATE}."
    )

    h1(doc, "12.  Contact Us")
    info_row(doc, "Address", ADDRESS)
    info_row(doc, "Telephone", PHONE)
    info_row(doc, "Email", EMAIL)
    info_row(doc, "Company No.", REG_NO)

    footer_note(doc)
    doc.save("Terms_and_Conditions.docx")
    print("✓  Terms_and_Conditions.docx")


# ── 3. Cookie Policy ──────────────────────────────────────────────────────────

def make_cookies():
    doc = new_doc("Cookie Policy")

    h1(doc, "1.  What Are Cookies?")
    body(doc,
        "Cookies are small text files placed on your device when you visit a website. They are widely used "
        "to make websites work more efficiently and to provide information to the website owner. We use "
        "cookies in accordance with the Privacy and Electronic Communications Regulations (PECR) and UK GDPR."
    )

    h1(doc, "2.  How We Use Cookies")
    body(doc, f"The {COMPANY} website ({WEBSITE}) uses the following categories of cookies:")

    h2(doc, "2.1  Strictly Necessary Cookies")
    body(doc,
        "These cookies are essential for the website to function and cannot be switched off. They are set "
        "in response to actions you take, such as filling in forms. They do not store personally identifiable information."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, t in enumerate(["Cookie name", "Purpose", "Duration"]):
        table.rows[0].cells[i].paragraphs[0].add_run(t).font.bold = True
    for row_data in [
        ("session_id", "Maintains your session as you navigate the website", "Session"),
        ("csrf_token", "Protects against cross-site request forgery", "Session"),
    ]:
        row = table.add_row().cells
        for i, v in enumerate(row_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()

    h2(doc, "2.2  Analytics Cookies")
    body(doc,
        "These cookies allow us to count visits and traffic sources so we can measure and improve the "
        "performance of our website. All information these cookies collect is aggregated and anonymous."
    )
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    for i, t in enumerate(["Cookie name", "Purpose", "Duration"]):
        table2.rows[0].cells[i].paragraphs[0].add_run(t).font.bold = True
    for row_data in [
        ("_ga", "Google Analytics — distinguishes users", "2 years"),
        ("_ga_*", "Google Analytics — maintains session state", "2 years"),
        ("_gid", "Google Analytics — distinguishes users", "24 hours"),
    ]:
        row = table2.add_row().cells
        for i, v in enumerate(row_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()

    h2(doc, "2.3  Functional Cookies")
    body(doc,
        "These cookies enable enhanced functionality and personalisation. They may be set by us or by "
        "third-party providers whose services we have added to our website."
    )

    h2(doc, "2.4  Marketing Cookies")
    body(doc,
        "We do not currently use marketing or advertising cookies on this website."
    )

    h1(doc, "3.  Your Cookie Choices")
    body(doc,
        "When you first visit our website, you will be presented with a cookie consent banner. "
        "You can choose to accept all cookies, or manage your preferences to accept only essential cookies. "
        "You can also manage cookies through your browser settings:"
    )
    for browser, path in [
        ("Google Chrome", "Settings > Privacy and security > Cookies and other site data"),
        ("Mozilla Firefox", "Options > Privacy & Security > Cookies and Site Data"),
        ("Safari", "Preferences > Privacy"),
        ("Microsoft Edge", "Settings > Cookies and site permissions"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{browser}: ").font.bold = True
        p.add_run(path)

    body(doc,
        "Please note that disabling certain cookies may affect the functionality of the website."
    )

    h1(doc, "4.  Third-Party Cookies")
    body(doc,
        "Some cookies on our website are placed by third-party services. These include Google Analytics "
        "(operated by Google LLC). These third parties have their own privacy policies and cookie policies "
        "which govern their use of such information."
    )

    h1(doc, "5.  More Information")
    body(doc, f"For more information about cookies, visit www.allaboutcookies.org.")
    body(doc,
        f"If you have any questions about our use of cookies, please contact us at {EMAIL} or "
        f"write to {ADDRESS}. This Cookie Policy was last updated on {DATE}."
    )

    footer_note(doc)
    doc.save("Cookie_Policy.docx")
    print("✓  Cookie_Policy.docx")


# ── 4. Complaints Procedure ───────────────────────────────────────────────────

def make_complaints():
    doc = new_doc("Complaints Procedure")

    h1(doc, "1.  Our Commitment")
    body(doc,
        f"{COMPANY} is committed to providing a high standard of service to all our clients. We take "
        f"all complaints seriously and aim to resolve them promptly, fairly, and professionally. "
        f"This procedure sets out how to raise a concern and what you can expect from us."
    )
    body(doc,
        f"{COMPANY} is a member of The Property Ombudsman (TPO) (Ref: {TPO_NO}), the government-approved "
        f"redress scheme for residential estate agents. If we are unable to resolve your complaint "
        f"internally, you have the right to refer the matter to TPO free of charge."
    )

    h1(doc, "2.  Stage 1 — Informal Resolution")
    body(doc,
        "In the first instance, please raise your concern with the member of staff dealing with your "
        "matter. Many concerns can be resolved quickly at this level without the need for a formal complaint."
    )
    body(doc,
        f"You can contact us by:"
    )
    for item in [
        f"Telephone: {PHONE} (Monday–Friday 9am–6pm, Saturday 10am–4pm)",
        f"Email: {EMAIL}",
        f"Post: {ADDRESS}",
    ]:
        bullet(doc, item)

    h1(doc, "3.  Stage 2 — Formal Written Complaint")
    body(doc,
        f"If your concern is not resolved informally, please submit a formal written complaint to the "
        f"Director at the address or email above, clearly marked 'Formal Complaint'. Please include:"
    )
    for item in [
        "Your full name and contact details",
        "The address of the property concerned",
        "A clear description of your complaint and what happened",
        "Copies of any relevant correspondence",
        "What outcome you are seeking",
    ]:
        bullet(doc, item)
    body(doc,
        "We will acknowledge your complaint in writing within 3 working days of receipt. "
        "We will then investigate fully and send you a written response within 15 working days. "
        "If we require more time, we will write to you to explain why and give a revised timescale."
    )

    h1(doc, "4.  Stage 3 — Director Review")
    body(doc,
        "If you remain dissatisfied with our Stage 2 response, you may request a review by writing "
        "to the Director within 14 days of receiving our Stage 2 response. The Director will review "
        "the matter independently and provide a Final Response within 15 working days."
    )

    h1(doc, "5.  The Property Ombudsman (TPO)")
    body(doc,
        "If you remain dissatisfied after receiving our Final Response — or if 8 weeks have elapsed "
        "since you raised your formal complaint — you may refer the matter to The Property Ombudsman:"
    )
    info_row(doc, "Website",   "www.tpos.co.uk")
    info_row(doc, "Telephone", "01722 333 306")
    info_row(doc, "Email",     "admin@tpos.co.uk")
    info_row(doc, "Post",      "The Property Ombudsman, Milford House, 43–55 Milford Street, Salisbury, SP1 2BP")
    body(doc,
        "The Property Ombudsman service is free of charge to consumers. You must refer your complaint "
        "to TPO within 12 months of our Final Response."
    )

    h1(doc, "6.  What TPO Can and Cannot Do")
    body(doc,
        "TPO can award compensation of up to £25,000 in certain circumstances. TPO considers complaints "
        "about conduct, not about the level of fees charged or commercial decisions we have made."
    )

    h1(doc, "7.  Keeping Records")
    body(doc,
        f"We keep records of all formal complaints and their outcomes for a minimum of 3 years in accordance "
        f"with our obligations as members of The Property Ombudsman scheme."
    )

    footer_note(doc)
    doc.save("Complaints_Procedure.docx")
    print("✓  Complaints_Procedure.docx")


# ── 5. AML Policy ────────────────────────────────────────────────────────────

def make_aml():
    doc = new_doc("Anti-Money Laundering Policy")

    h1(doc, "1.  Introduction and Legal Framework")
    body(doc,
        f"{COMPANY} is registered with HMRC for Anti-Money Laundering supervision under the "
        f"Money Laundering, Terrorist Financing and Transfer of Funds (Information on the Payer) "
        f"Regulations 2017 (the 'MLRs 2017') (HMRC Registration: {HMRC_MLR}). "
        f"We are also subject to the Proceeds of Crime Act 2002 and the Terrorism Act 2000."
    )
    body(doc,
        "All partners and staff involved in estate agency work are aware of their obligations under "
        "the MLRs 2017 and have received appropriate training."
    )

    h1(doc, "2.  Nominated Officer")
    body(doc,
        f"Our Money Laundering Reporting Officer (MLRO) is Max Watson, Director. "
        f"All internal suspicious activity reports should be submitted to the MLRO in writing. "
        f"Contact: {EMAIL}"
    )

    h1(doc, "3.  Customer Due Diligence (CDD)")
    h2(doc, "3.1  When CDD is required")
    body(doc, "We carry out Customer Due Diligence checks before establishing a business relationship with:")
    for item in [
        "Vendors (sellers) instructing us to market their property",
        "Purchasers making an offer on a property",
        "Landlords instructing us to let their property",
        "Tenants entering a tenancy agreement",
    ]:
        bullet(doc, item)

    h2(doc, "3.2  Standard CDD measures")
    body(doc, "For individuals, we will request:")
    for item in [
        "One document confirming identity: valid passport, UK driving licence, or national identity card",
        "One document confirming address (dated within 3 months): utility bill, bank statement, or official government letter",
        "Source of funds / source of wealth documentation for cash buyers or where transaction value is £10,000 or more",
    ]:
        bullet(doc, item)
    body(doc, "For corporate clients (companies, LLPs, trusts), we will request:")
    for item in [
        "Certificate of incorporation and articles of association",
        "Confirmation of beneficial owners (persons with 25% or more ownership)",
        "Identity verification for each beneficial owner",
        "Evidence of the authority of the individual instructing us",
    ]:
        bullet(doc, item)

    h2(doc, "3.3  Enhanced Due Diligence")
    body(doc,
        "Enhanced Due Diligence (EDD) is applied in the following higher-risk situations:"
    )
    for item in [
        "Transactions involving Politically Exposed Persons (PEPs) or their associates",
        "Transactions involving high-risk third countries designated by the FATF",
        "Complex or unusually large transactions with no apparent economic purpose",
        "Cash-heavy transactions or unusual payment structures",
    ]:
        bullet(doc, item)

    h2(doc, "3.4  Ongoing monitoring")
    body(doc,
        "We monitor all transactions and business relationships on an ongoing basis. We will request "
        "updated verification documents if there are changes in circumstances or if our records are out of date."
    )

    h1(doc, "4.  Record Keeping")
    body(doc,
        "We retain copies of all CDD documentation for a minimum of 5 years from the end of the "
        "business relationship, in accordance with Regulation 40 of the MLRs 2017."
    )

    h1(doc, "5.  Suspicious Activity Reporting")
    body(doc,
        "Where any member of staff knows or suspects that a client is engaged in money laundering "
        "or terrorist financing, they must immediately report this to the MLRO by completing an "
        "internal Suspicious Activity Report (SAR). The MLRO will then determine whether to submit "
        "a SAR to the National Crime Agency (NCA) via the NCA's online portal."
    )
    body(doc,
        "IMPORTANT: Staff must not 'tip off' a client that a SAR has been filed or that they are "
        "under investigation. Tipping off is a criminal offence under Section 333A of POCA 2002."
    )

    h1(doc, "6.  Training")
    body(doc,
        "All relevant staff receive AML training upon joining the firm and at regular intervals "
        "thereafter, and whenever there are material changes to our obligations or procedures. "
        "Training records are maintained by the MLRO."
    )

    h1(doc, "7.  Policy Review")
    body(doc,
        f"This policy is reviewed annually by the MLRO and updated as necessary to reflect "
        f"changes in legislation, guidance, and best practice. Last reviewed: {DATE}."
    )

    footer_note(doc)
    doc.save("Anti_Money_Laundering_Policy.docx")
    print("✓  Anti_Money_Laundering_Policy.docx")


# ── 6. Fees Schedule ──────────────────────────────────────────────────────────

def make_fees():
    doc = new_doc("Fees & Charges Schedule")

    h1(doc, "1.  Introduction")
    body(doc,
        f"In accordance with the Consumer Protection from Unfair Trading Regulations 2008 and "
        f"our obligations as a member of The Property Ombudsman, {COMPANY} is committed to "
        f"full transparency regarding all fees and charges. This schedule sets out the fees that "
        f"may apply to our services."
    )
    body(doc,
        "All fees are exclusive of VAT unless otherwise stated. VAT at the current rate will be "
        "applied to all fees. Specific fee arrangements will be confirmed in your Letter of Engagement "
        "or Lettings Agreement before you commit to using our services."
    )

    h1(doc, "2.  Sales Fees")
    h2(doc, "2.1  Sole agency")
    body(doc,
        "Our standard commission for selling your property on a sole agency basis is 1.25% + VAT "
        "of the agreed sale price (minimum fee £3,000 + VAT). This fee is payable upon exchange of contracts."
    )
    h2(doc, "2.2  Joint agency")
    body(doc,
        "Where you instruct us alongside another agent, our fee is 1.75% + VAT of the agreed sale price. "
        "In a joint agency arrangement the full fee is payable to whichever agent introduces the buyer."
    )
    h2(doc, "2.3  Additional sales services")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for i, t in enumerate(["Service", "Fee"]):
        table.rows[0].cells[i].paragraphs[0].add_run(t).font.bold = True
    for row_data in [
        ("Professional photography (included in sole agency)", "Included"),
        ("Floorplan (included in sole agency)", "Included"),
        ("Energy Performance Certificate (EPC)", "£90 + VAT"),
        ("Enhanced photography / drone photography", "From £250 + VAT"),
        ("Accompanied viewings (if required beyond standard service)", "£75 per viewing + VAT"),
    ]:
        row = table.add_row().cells
        for i, v in enumerate(row_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()

    h1(doc, "3.  Lettings Fees — Landlords")
    h2(doc, "3.1  Tenant-find service")
    body(doc,
        "We will market your property, conduct viewings, reference tenants, and prepare a tenancy agreement. "
        "Fee: equivalent to one month's rent + VAT (minimum £750 + VAT)."
    )
    h2(doc, "3.2  Fully managed service")
    body(doc,
        "In addition to the tenant-find service, we collect rent, manage maintenance and repairs, "
        "conduct periodic inspections, and handle all tenant communications throughout the tenancy. "
        "Fee: 12% of the monthly rent + VAT (ongoing, collected monthly)."
    )
    h2(doc, "3.3  Additional landlord services")
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    for i, t in enumerate(["Service", "Fee"]):
        table2.rows[0].cells[i].paragraphs[0].add_run(t).font.bold = True
    for row_data in [
        ("Inventory preparation (check-in)", "From £120 + VAT"),
        ("Inventory check-out report", "From £100 + VAT"),
        ("Deposit registration and protection", "£50 + VAT"),
        ("Tenancy renewal (fixed term)", "£150 + VAT"),
        ("Section 21 / Section 8 notice service", "£200 + VAT"),
        ("Energy Performance Certificate (EPC)", "£90 + VAT"),
        ("Gas Safety Certificate (arrangement)", "At cost"),
        ("EICR (Electrical Installation Condition Report)", "At cost"),
    ]:
        row = table2.add_row().cells
        for i, v in enumerate(row_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()

    h1(doc, "4.  Fees Payable by Tenants")
    body(doc,
        "In accordance with the Tenant Fees Act 2019, the only payments that can be charged to "
        "tenants in connection with a residential tenancy in England are:"
    )
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    for i, t in enumerate(["Permitted payment", "Amount / Notes"]):
        table3.rows[0].cells[i].paragraphs[0].add_run(t).font.bold = True
    for row_data in [
        ("Rent", "As agreed in the tenancy agreement"),
        ("Refundable tenancy deposit", "Maximum 5 weeks' rent (annual rent under £50,000)"),
        ("Refundable holding deposit", "Maximum 1 week's rent; returned within 15 days"),
        ("Changes to the tenancy (tenant request)", "Maximum £50 + VAT per variation"),
        ("Early termination (tenant request)", "Landlord's reasonable costs only"),
        ("Utilities and council tax", "As applicable and set out in the tenancy agreement"),
        ("Default charges — late rent", "Interest at 3% above Bank of England base rate (after 14 days)"),
        ("Default charges — lost key/security device", "Reasonable costs evidenced in writing"),
    ]:
        row = table3.add_row().cells
        for i, v in enumerate(row_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()
    body(doc,
        "We do not charge any fees to tenants that are not permitted under the Tenant Fees Act 2019. "
        "Any prohibited payment will be refunded in full."
    )

    h1(doc, "5.  Valuation Services")
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    for i, t in enumerate(["Service", "Fee"]):
        table4.rows[0].cells[i].paragraphs[0].add_run(t).font.bold = True
    for row_data in [
        ("Market appraisal (for sale or let)", "Free, no obligation"),
        ("Written valuation for probate / matrimonial / other legal purposes", "From £250 + VAT"),
    ]:
        row = table4.add_row().cells
        for i, v in enumerate(row_data):
            row[i].paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()

    h1(doc, "6.  Client Money Protection")
    body(doc,
        f"Where we handle client money, this is protected by the {CMP_PROV} Client Money Protection "
        f"scheme. A copy of our CMP certificate is available upon request and is displayed in our office."
    )

    h1(doc, "7.  Redress Scheme")
    body(doc,
        f"{COMPANY} is a member of The Property Ombudsman (TPO) redress scheme (Ref: {TPO_NO}). "
        f"TPO membership details are displayed in our office and on our website."
    )

    h1(doc, "8.  Changes to Fees")
    body(doc,
        f"We reserve the right to amend our fees from time to time. Any changes will be notified "
        f"to you in writing before they take effect. This schedule was last updated on {DATE}."
    )

    footer_note(doc)
    doc.save("Fees_and_Charges_Schedule.docx")
    print("✓  Fees_and_Charges_Schedule.docx")


# ── 7. Accessibility Statement ────────────────────────────────────────────────

def make_accessibility():
    doc = new_doc("Accessibility Statement")

    h1(doc, "1.  Our Commitment")
    body(doc,
        f"{COMPANY} is committed to making {WEBSITE} accessible to all users, including those with "
        f"disabilities. We continually seek to improve the accessibility of our website in accordance "
        f"with the Web Content Accessibility Guidelines (WCAG) 2.1."
    )

    h1(doc, "2.  Conformance Status")
    body(doc,
        f"We aim to meet WCAG 2.1 Level AA conformance. We are working towards full compliance "
        f"and are aware of some areas where improvements are needed (see Section 4 below)."
    )

    h1(doc, "3.  Technical Specification")
    body(doc,
        "Our website relies on the following technologies to work with your browser and any assistive technology you use:"
    )
    for item in ["HTML5", "CSS3", "JavaScript", "WAI-ARIA where appropriate"]:
        bullet(doc, item)

    h1(doc, "4.  Known Accessibility Issues")
    body(doc,
        "We are aware of the following limitations and are working to address them:"
    )
    for item in [
        "Some older images may lack descriptive alt text — we are reviewing all images",
        "Some PDF documents linked from the website may not be fully screen-reader accessible — we aim to provide accessible HTML alternatives",
        "Some interactive elements may not be fully keyboard-navigable — we are working to improve this",
    ]:
        bullet(doc, item)

    h1(doc, "5.  Assistive Technology Compatibility")
    body(doc,
        "Our website is designed to be compatible with the following assistive technologies:"
    )
    for item in [
        "Screen readers (including NVDA, JAWS, VoiceOver)",
        "Speech recognition software",
        "Screen magnification software",
        "Keyboard-only navigation",
    ]:
        bullet(doc, item)

    h1(doc, "6.  Feedback and Contact")
    body(doc,
        "If you experience any accessibility barriers on our website, or if you require information "
        "in an alternative format, please contact us:"
    )
    info_row(doc, "Email", EMAIL)
    info_row(doc, "Telephone", PHONE)
    info_row(doc, "Post", ADDRESS)
    body(doc, "We aim to respond to accessibility feedback within 5 working days.")

    h1(doc, "7.  Enforcement")
    body(doc,
        "If you are not satisfied with our response, you may contact the Equality Advisory and Support "
        "Service (EASS): www.equalityadvisoryservice.com"
    )

    h1(doc, "8.  Review")
    body(doc,
        f"This accessibility statement was prepared on {DATE} and will be reviewed annually. "
        f"We welcome feedback to help us improve the accessibility of our website."
    )

    footer_note(doc)
    doc.save("Accessibility_Statement.docx")
    print("✓  Accessibility_Statement.docx")


# ── Run all ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"\nGenerating {COMPANY} legal documents…\n")
    make_privacy_policy()
    make_terms()
    make_cookies()
    make_complaints()
    make_aml()
    make_fees()
    make_accessibility()
    print(f"\nAll documents saved to: {__import__('os').getcwd()}\n")
